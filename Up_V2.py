"""
libby_up.py  —  Up: Document Mirror Engine for Libby
=====================================================
Mirrors a source folder into a clean library of .txt files
that Libby reads from. Originals are never touched.

Mirror layout (relative to source folder):
  source/
    _libby_mirror/
      <mirrored .txt files — preserving subfolder structure>
    _libby_mirror/_archive/
      <.txt files whose originals were deleted>

Supported source types: .txt  .pdf  .xlsx  .docx
"""

import os
import shutil
import hashlib
import sqlite3
import threading
import json
from datetime import datetime
from pathlib import Path

import fitz                          # pymupdf  — PDF
import openpyxl                      # xlsx
from docx import Document            # python-docx

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
MIRROR_DIR   = "_libby_mirror"
ARCHIVE_DIR  = "_archive"
DB_NAME      = "up_index.db"
SUPPORTED    = {".txt", ".pdf", ".xlsx", ".docx"}

# ─────────────────────────────────────────────
# PATHS
# ─────────────────────────────────────────────

def mirror_root(source_folder: str) -> str:
    return os.path.join(source_folder, MIRROR_DIR)

def archive_root(source_folder: str) -> str:
    return os.path.join(mirror_root(source_folder), ARCHIVE_DIR)

def db_path(source_folder: str) -> str:
    return os.path.join(mirror_root(source_folder), DB_NAME)

def mirrored_path(source_folder: str, source_file: str) -> str:
    """Return the .txt path in the mirror that corresponds to source_file."""
    rel   = os.path.relpath(source_file, source_folder)
    parts = list(Path(rel).parts)
    # Drop the filename, replace extension with .txt
    stem  = Path(parts[-1]).stem
    parts[-1] = stem + ".txt"
    return os.path.join(mirror_root(source_folder), *parts)

# ─────────────────────────────────────────────
# SQLITE — INDEX
# ─────────────────────────────────────────────

def open_db(source_folder: str) -> sqlite3.Connection:
    os.makedirs(mirror_root(source_folder), exist_ok=True)
    con = sqlite3.connect(db_path(source_folder), check_same_thread=False)
    con.row_factory = sqlite3.Row
    _create_schema(con)
    return con

def _create_schema(con: sqlite3.Connection):
    con.executescript("""
        -- ── LEVEL 0: Documents ──────────────────────────────────────────────
        -- One row per source file. Tracks sync state and structured metadata.
        CREATE TABLE IF NOT EXISTS documents (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            source_path   TEXT    UNIQUE NOT NULL,  -- absolute path to original
            mirror_path   TEXT    NOT NULL,          -- absolute path to .txt copy
            file_type     TEXT    NOT NULL,          -- txt | pdf | xlsx | docx
            file_hash     TEXT    NOT NULL,          -- SHA-256, change detection
            status        TEXT    NOT NULL DEFAULT 'active', -- active | archived | error
            first_seen    TEXT    NOT NULL,          -- ISO timestamp
            last_synced   TEXT    NOT NULL,          -- ISO timestamp
            char_count    INTEGER DEFAULT 0,
            -- Structured metadata (filterable)
            author        TEXT,
            doc_date      TEXT,                      -- ISO date if detectable
            subfolder     TEXT,                      -- relative subfolder path
            -- Free-form tags stored as JSON array e.g. '["finance","Q3"]'
            tags          TEXT    NOT NULL DEFAULT '[]',
            error         TEXT
        );

        -- ── LEVEL 1: Sections ───────────────────────────────────────────────
        -- Logical divisions within a document (pages, sheets, headings).
        -- For PDFs: one section per page.
        -- For XLSX: one section per sheet.
        -- For DOCX/TXT: one section per detected heading or every N paragraphs.
        CREATE TABLE IF NOT EXISTS sections (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            document_id   INTEGER NOT NULL REFERENCES documents(id) ON DELETE CASCADE,
            section_index INTEGER NOT NULL,   -- 0-based position in document
            section_type  TEXT    NOT NULL,   -- page | sheet | heading | block
            heading       TEXT,               -- heading text if detected
            char_count    INTEGER DEFAULT 0
        );

        -- ── LEVEL 2: Chunks ─────────────────────────────────────────────────
        -- Fixed-size overlapping text windows within a section.
        -- These are what ChromaDB will embed for semantic search.
        CREATE TABLE IF NOT EXISTS chunks (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            document_id   INTEGER NOT NULL REFERENCES documents(id) ON DELETE CASCADE,
            section_id    INTEGER NOT NULL REFERENCES sections(id)  ON DELETE CASCADE,
            chunk_index   INTEGER NOT NULL,   -- 0-based position within section
            text          TEXT    NOT NULL,
            char_start    INTEGER,            -- char offset from section start
            char_end      INTEGER,
            char_count    INTEGER DEFAULT 0,
            chroma_id     TEXT                -- ChromaDB vector ID (set when embedded)
        );

        -- ── LEVEL 3: Sentences ──────────────────────────────────────────────
        -- Individual sentences within a chunk.
        -- Enables precise positional queries: "last sentence in newest file".
        CREATE TABLE IF NOT EXISTS sentences (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            document_id     INTEGER NOT NULL REFERENCES documents(id) ON DELETE CASCADE,
            chunk_id        INTEGER NOT NULL REFERENCES chunks(id)    ON DELETE CASCADE,
            sentence_index  INTEGER NOT NULL,   -- 0-based position within chunk
            text            TEXT    NOT NULL,
            char_count      INTEGER DEFAULT 0
        );

        -- ── TAGS ────────────────────────────────────────────────────────────
        -- Normalised tag table for efficient tag-based queries.
        -- Mirrors the JSON tags array on documents for JOIN-based filtering.
        CREATE TABLE IF NOT EXISTS document_tags (
            document_id  INTEGER NOT NULL REFERENCES documents(id) ON DELETE CASCADE,
            tag          TEXT    NOT NULL,
            PRIMARY KEY (document_id, tag)
        );

        -- ── SYNC LOG ────────────────────────────────────────────────────────
        CREATE TABLE IF NOT EXISTS sync_log (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            event       TEXT NOT NULL,   -- added | updated | archived | error
            source_path TEXT NOT NULL,
            detail      TEXT,
            ts          TEXT NOT NULL
        );

        -- ── INDEXES ─────────────────────────────────────────────────────────
        CREATE INDEX IF NOT EXISTS idx_sections_document  ON sections(document_id);
        CREATE INDEX IF NOT EXISTS idx_chunks_document    ON chunks(document_id);
        CREATE INDEX IF NOT EXISTS idx_chunks_section     ON chunks(section_id);
        CREATE INDEX IF NOT EXISTS idx_sentences_chunk    ON sentences(chunk_id);
        CREATE INDEX IF NOT EXISTS idx_sentences_document ON sentences(document_id);
        CREATE INDEX IF NOT EXISTS idx_tags_tag           ON document_tags(tag);
        CREATE INDEX IF NOT EXISTS idx_docs_subfolder     ON documents(subfolder);
        CREATE INDEX IF NOT EXISTS idx_docs_file_type     ON documents(file_type);
        CREATE INDEX IF NOT EXISTS idx_docs_status        ON documents(status);
    """)
    con.commit()

def _log_event(con, event: str, source_path: str, detail: str = ""):
    con.execute(
        "INSERT INTO sync_log (event, source_path, detail, ts) VALUES (?,?,?,?)",
        (event, source_path, detail, datetime.now().isoformat())
    )
    con.commit()

# ─────────────────────────────────────────────
# HASHING
# ─────────────────────────────────────────────

def file_hash(filepath: str) -> str:
    """SHA-256 of the file — used to detect real changes."""
    h = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()

# ─────────────────────────────────────────────
# READERS
# ─────────────────────────────────────────────

def read_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()

def read_pdf(filepath: str) -> str:
    doc = fitz.open(filepath)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text().strip()
        if text:
            pages.append(f"[Page {i+1}]\n{text}")
    doc.close()
    return "\n\n".join(pages)

def read_xlsx(filepath: str) -> str:
    wb = openpyxl.load_workbook(filepath, data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws  = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cleaned = [str(c).strip() if c is not None else "" for c in row]
            if any(cleaned):
                rows.append(" | ".join(cleaned))
        if rows:
            sheets.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets)

def read_docx(filepath: str) -> str:
    doc    = Document(filepath)
    paras  = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Also capture tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n\n".join(paras)

READERS = {
    ".txt":  read_txt,
    ".pdf":  read_pdf,
    ".xlsx": read_xlsx,
    ".docx": read_docx,
}

def extract_text(filepath: str) -> tuple[str, str | None]:
    """
    Returns (text, error).
    error is None on success, a string on failure.
    """
    ext = Path(filepath).suffix.lower()
    reader = READERS.get(ext)
    if not reader:
        return "", f"Unsupported type: {ext}"
    try:
        text = reader(filepath)
        return text.strip(), None
    except Exception as e:
        return "", str(e)

# ─────────────────────────────────────────────
# WRITE MIRROR FILE
# ─────────────────────────────────────────────

def write_mirror(source_folder: str, source_file: str, text: str) -> str:
    """Write text to the appropriate .txt path in the mirror. Returns mirror path."""
    dest = mirrored_path(source_folder, source_file)
    os.makedirs(os.path.dirname(dest), exist_ok=True)
    with open(dest, "w", encoding="utf-8") as f:
        # Header so Libby knows provenance
        rel = os.path.relpath(source_file, source_folder)
        f.write(f"[Source: {rel}]\n")
        f.write(f"[Converted: {datetime.now().strftime('%Y-%m-%d %H:%M')}]\n\n")
        f.write(text)
    return dest

# ─────────────────────────────────────────────
# METADATA INDEXING
# ─────────────────────────────────────────────
# Populates sections → chunks → sentences for a
# document after its .txt mirror has been written.
# Wipes and rebuilds on every update (hash changed).

CHUNK_SIZE    = 1200   # chars per chunk
CHUNK_OVERLAP = 150    # overlap between chunks


def _split_sentences(text: str) -> list[str]:
    """
    Naive sentence splitter — no NLTK dependency.
    Splits on '. ', '? ', '! ' and newlines.
    """
    import re
    parts = re.split(r'(?<=[.?!])\s+|\n', text)
    return [p.strip() for p in parts if p.strip()]


def _chunk_text(text: str) -> list[dict]:
    """
    Split text into overlapping chunks.
    Returns list of {text, char_start, char_end}.
    """
    chunks = []
    start  = 0
    while start < len(text):
        end  = min(start + CHUNK_SIZE, len(text))
        chunk_text = text[start:end].strip()
        if len(chunk_text) > 50:
            chunks.append({
                "text":       chunk_text,
                "char_start": start,
                "char_end":   end,
            })
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


def _detect_sections_txt(text: str) -> list[dict]:
    """
    For .txt and .docx: split on blank lines or heading-like lines.
    Returns list of {heading, text, section_type}.
    """
    import re
    blocks  = re.split(r'\n{2,}', text)
    sections = []
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines   = block.splitlines()
        heading = None
        # A heading: first line is short, title-case or all-caps, no period
        first = lines[0].strip()
        if len(first) < 80 and not first.endswith('.') and (
            first.istitle() or first.isupper() or first.startswith('#')
        ):
            heading = first.lstrip('#').strip()
        sections.append({
            "heading":      heading,
            "text":         block,
            "section_type": "heading" if heading else "block",
        })
    return sections or [{"heading": None, "text": text, "section_type": "block"}]


def _detect_sections_pdf(text: str) -> list[dict]:
    """For PDFs: each [Page N] block is a section."""
    import re
    pages    = re.split(r'\[Page \d+\]', text)
    page_nos = re.findall(r'\[Page (\d+)\]', text)
    sections = []
    for i, content in enumerate(pages):
        content = content.strip()
        if not content:
            continue
        pn = page_nos[i - 1] if i > 0 and i - 1 < len(page_nos) else str(i)
        sections.append({
            "heading":      f"Page {pn}",
            "text":         content,
            "section_type": "page",
        })
    return sections or [{"heading": None, "text": text, "section_type": "page"}]


def _detect_sections_xlsx(text: str) -> list[dict]:
    """For XLSX: each [Sheet: name] block is a section."""
    import re
    sheets    = re.split(r'\[Sheet: [^\]]+\]', text)
    names     = re.findall(r'\[Sheet: ([^\]]+)\]', text)
    sections  = []
    for i, content in enumerate(sheets):
        content = content.strip()
        if not content:
            continue
        name = names[i - 1] if i > 0 and i - 1 < len(names) else f"Sheet {i}"
        sections.append({
            "heading":      name,
            "text":         content,
            "section_type": "sheet",
        })
    return sections or [{"heading": None, "text": text, "section_type": "sheet"}]


SECTION_DETECTORS = {
    ".txt":  _detect_sections_txt,
    ".docx": _detect_sections_txt,
    ".pdf":  _detect_sections_pdf,
    ".xlsx": _detect_sections_xlsx,
}


def index_document(con: sqlite3.Connection,
                   document_id: int,
                   text: str,
                   file_type: str):
    """
    Wipe and rebuild sections → chunks → sentences for document_id.
    Called after every successful convert (add or update).
    """
    # Wipe existing hierarchy for this document
    con.execute("DELETE FROM sentences WHERE document_id=?", (document_id,))
    con.execute("DELETE FROM chunks    WHERE document_id=?", (document_id,))
    con.execute("DELETE FROM sections  WHERE document_id=?", (document_id,))
    con.commit()

    detector = SECTION_DETECTORS.get(file_type, _detect_sections_txt)
    raw_sections = detector(text)

    for sec_idx, sec in enumerate(raw_sections):
        # ── Insert section ──
        cur = con.execute(
            """INSERT INTO sections
               (document_id, section_index, section_type, heading, char_count)
               VALUES (?,?,?,?,?)""",
            (document_id, sec_idx, sec["section_type"],
             sec.get("heading"), len(sec["text"]))
        )
        section_id = cur.lastrowid

        raw_chunks = _chunk_text(sec["text"])

        for chk_idx, chk in enumerate(raw_chunks):
            # ── Insert chunk ──
            cur2 = con.execute(
                """INSERT INTO chunks
                   (document_id, section_id, chunk_index,
                    text, char_start, char_end, char_count)
                   VALUES (?,?,?,?,?,?,?)""",
                (document_id, section_id, chk_idx,
                 chk["text"], chk["char_start"], chk["char_end"],
                 len(chk["text"]))
            )
            chunk_id = cur2.lastrowid

            sentences = _split_sentences(chk["text"])

            for sent_idx, sent in enumerate(sentences):
                con.execute(
                    """INSERT INTO sentences
                       (document_id, chunk_id, sentence_index, text, char_count)
                       VALUES (?,?,?,?,?)""",
                    (document_id, chunk_id, sent_idx, sent, len(sent))
                )

    con.commit()


def set_document_tags(con: sqlite3.Connection,
                      document_id: int,
                      tags: list[str]):
    """
    Replace all tags for a document.
    Tags are stored both normalised (document_tags) and as JSON on documents.
    """
    import json as _json
    con.execute("DELETE FROM document_tags WHERE document_id=?", (document_id,))
    for tag in tags:
        tag = tag.strip().lower()
        if tag:
            con.execute(
                "INSERT OR IGNORE INTO document_tags (document_id, tag) VALUES (?,?)",
                (document_id, tag)
            )
    con.execute(
        "UPDATE documents SET tags=? WHERE id=?",
        (_json.dumps(tags), document_id)
    )
    con.commit()


def infer_tags(source_path: str, subfolder: str) -> list[str]:
    """
    Rule-based tag inference from filename and subfolder.
    No LLM required — pure heuristics for now.
    """
    tags     = []
    combined = (os.path.basename(source_path) + " " + (subfolder or "")).lower()

    tag_rules = {
        "finance":   ["invoice","budget","finance","financial","revenue","cost","expense","p&l","profit"],
        "sales":     ["sales","order","quote","pricing","customer","crm"],
        "inventory": ["inventory","stock","warehouse","sku","reorder"],
        "hr":        ["employee","staff","salary","hr","payroll","leave"],
        "legal":     ["contract","agreement","legal","compliance","policy","nda"],
        "report":    ["report","summary","analysis","review","monthly","quarterly","annual"],
        "guide":     ["guide","manual","handbook","faq","howto","instructions"],
        "project":   ["project","plan","timeline","milestone","roadmap"],
    }

    for tag, keywords in tag_rules.items():
        if any(kw in combined for kw in keywords):
            tags.append(tag)

    # Add subfolder as a tag if it's meaningful
    if subfolder and subfolder not in (".", ""):
        folder_tag = subfolder.strip("/\\").replace(" ", "_").lower()
        if folder_tag not in tags:
            tags.append(folder_tag)

    return tags


# ─────────────────────────────────────────────
# ARCHIVE
# ─────────────────────────────────────────────

def archive_mirror(source_folder: str, mirror_file: str):
    """Move a .txt mirror file to _archive/ when its source is deleted."""
    arc = archive_root(source_folder)
    os.makedirs(arc, exist_ok=True)
    basename = os.path.basename(mirror_file)
    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    dest     = os.path.join(arc, f"{Path(basename).stem}_{ts}.txt")
    if os.path.exists(mirror_file):
        shutil.move(mirror_file, dest)
    return dest

# ─────────────────────────────────────────────
# CORE SYNC ENGINE
# ─────────────────────────────────────────────

class SyncResult:
    def __init__(self):
        self.added    = []   # newly converted files
        self.updated  = []   # re-converted (changed source)
        self.archived = []   # source deleted → mirror archived
        self.skipped  = []   # unchanged
        self.errors   = []   # (filepath, error_message)

    def summary(self) -> str:
        lines = [
            f"✅  Added:    {len(self.added)}",
            f"🔄  Updated:  {len(self.updated)}",
            f"📦  Archived: {len(self.archived)}",
            f"⏭   Skipped:  {len(self.skipped)}",
        ]
        if self.errors:
            lines.append(f"⚠   Errors:   {len(self.errors)}")
        return "\n".join(lines)


def sync(source_folder: str,
         on_progress=None,
         con: sqlite3.Connection = None) -> SyncResult:
    """
    Full sync of source_folder → mirror.

    on_progress(event, filepath) is called for each action.
    Pass an existing DB connection or one will be opened.
    """
    result      = SyncResult()
    owns_con    = con is None
    if owns_con:
        con = open_db(source_folder)

    mirror = mirror_root(source_folder)
    arc    = archive_root(source_folder)

    # ── Discover all source files ──────────────────────────────────────
    source_files = {}
    for root, dirs, files in os.walk(source_folder):
        # Skip the mirror folder itself
        dirs[:] = [d for d in dirs if os.path.join(root, d) != mirror]
        for fname in files:
            ext = Path(fname).suffix.lower()
            if ext not in SUPPORTED:
                continue
            fp = os.path.join(root, fname)
            source_files[fp] = ext

    now = datetime.now().isoformat()

    # ── Process each source file ────────────────────────────────────────
    for filepath, ext in source_files.items():
        row = con.execute(
            "SELECT * FROM documents WHERE source_path=?", (filepath,)
        ).fetchone()

        try:
            fh = file_hash(filepath)
        except Exception as e:
            result.errors.append((filepath, f"Hash error: {e}"))
            continue

        if row is None:
            # New file — convert and index
            text, err = extract_text(filepath)
            if err:
                result.errors.append((filepath, err))
                _log_event(con, "error", filepath, err)
                con.execute(
                    """INSERT INTO documents
                       (source_path, mirror_path, file_type, file_hash,
                        status, first_seen, last_synced, char_count, error)
                       VALUES (?,?,?,?,?,?,?,?,?)""",
                    (filepath, "", ext, fh, "error", now, now, 0, err)
                )
                con.commit()
                continue

            mirror_fp  = write_mirror(source_folder, filepath, text)
            subfolder  = os.path.relpath(os.path.dirname(filepath), source_folder)
            cur = con.execute(
                """INSERT INTO documents
                   (source_path, mirror_path, file_type, file_hash,
                    status, first_seen, last_synced, char_count, subfolder)
                   VALUES (?,?,?,?,?,?,?,?,?)""",
                (filepath, mirror_fp, ext, fh, "active", now, now,
                 len(text), subfolder)
            )
            doc_id = cur.lastrowid
            con.commit()
            # Build section → chunk → sentence hierarchy
            index_document(con, doc_id, text, ext)
            # Infer and store tags
            tags = infer_tags(filepath, subfolder)
            set_document_tags(con, doc_id, tags)
            _log_event(con, "added", filepath)
            result.added.append(filepath)
            if on_progress:
                on_progress("added", filepath)

        elif row["file_hash"] != fh:
            # Changed file — re-convert and overwrite mirror
            text, err = extract_text(filepath)
            if err:
                result.errors.append((filepath, err))
                _log_event(con, "error", filepath, err)
                continue

            mirror_fp = write_mirror(source_folder, filepath, text)
            con.execute(
                """UPDATE documents SET
                   file_hash=?, mirror_path=?, status='active',
                   last_synced=?, char_count=?, error=NULL
                   WHERE source_path=?""",
                (fh, mirror_fp, now, len(text), filepath)
            )
            doc_id = con.execute(
                "SELECT id FROM documents WHERE source_path=?", (filepath,)
            ).fetchone()["id"]
            con.commit()
            # Rebuild hierarchy for changed file
            index_document(con, doc_id, text, ext)
            subfolder = os.path.relpath(os.path.dirname(filepath), source_folder)
            tags = infer_tags(filepath, subfolder)
            set_document_tags(con, doc_id, tags)
            _log_event(con, "updated", filepath)
            result.updated.append(filepath)
            if on_progress:
                on_progress("updated", filepath)

        else:
            # Unchanged
            result.skipped.append(filepath)

    # ── Detect deletions ────────────────────────────────────────────────
    indexed = con.execute(
        "SELECT source_path, mirror_path FROM documents WHERE status='active'"
    ).fetchall()

    for row in indexed:
        sp = row["source_path"]
        mp = row["mirror_path"]
        if sp not in source_files and os.path.exists(sp) is False:
            # Source gone — archive the mirror
            arc_dest = archive_mirror(source_folder, mp)
            con.execute(
                """UPDATE documents SET status='archived', last_synced=?
                   WHERE source_path=?""",
                (now, sp)
            )
            con.commit()
            _log_event(con, "archived", sp, f"→ {arc_dest}")
            result.archived.append(sp)
            if on_progress:
                on_progress("archived", sp)

    if owns_con:
        con.close()

    return result


# ─────────────────────────────────────────────
# THREADED SYNC (for UI use)
# ─────────────────────────────────────────────

def sync_async(source_folder: str,
               on_progress=None,
               on_complete=None,
               con: sqlite3.Connection = None):
    """
    Run sync() on a daemon thread.
    on_complete(result: SyncResult) is called on the thread when done.
    """
    def _run():
        result = sync(source_folder, on_progress=on_progress, con=con)
        if on_complete:
            on_complete(result)

    t = threading.Thread(target=_run, daemon=True)
    t.start()
    return t


# ─────────────────────────────────────────────
# STATUS QUERY
# ─────────────────────────────────────────────

def get_status(source_folder: str) -> dict:
    """
    Returns a dict summary of the current mirror state.
    Safe to call even if the DB doesn't exist yet.
    """
    if not os.path.exists(db_path(source_folder)):
        return {"indexed": 0, "archived": 0, "errors": 0,
                "sections": 0, "chunks": 0, "sentences": 0,
                "never_synced": True}
    con = open_db(source_folder)
    doc_rows = con.execute(
        "SELECT status, COUNT(*) as n FROM documents GROUP BY status"
    ).fetchall()
    counts = {r["status"]: r["n"] for r in doc_rows}
    sections  = con.execute("SELECT COUNT(*) FROM sections").fetchone()[0]
    chunks    = con.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
    sentences = con.execute("SELECT COUNT(*) FROM sentences").fetchone()[0]
    con.close()
    return {
        "indexed":      counts.get("active",   0),
        "archived":     counts.get("archived", 0),
        "errors":       counts.get("error",    0),
        "sections":     sections,
        "chunks":       chunks,
        "sentences":    sentences,
        "never_synced": False,
    }


# ─────────────────────────────────────────────
# CLI ENTRY POINT  (python libby_up.py <folder>)
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python libby_up.py <source_folder>")
        sys.exit(1)

    folder = sys.argv[1]
    if not os.path.isdir(folder):
        print(f"Error: not a directory: {folder}")
        sys.exit(1)

    print(f"Up  —  syncing: {folder}")
    print(f"Mirror: {mirror_root(folder)}\n")

    def progress(event, filepath):
        icons = {"added": "➕", "updated": "🔄", "archived": "📦"}
        print(f"  {icons.get(event,'·')}  [{event}]  {os.path.basename(filepath)}")

    result = sync(folder, on_progress=progress)
    print(f"\n{result.summary()}")

    # Show metadata counts from DB
    status = get_status(folder)
    print(f"\nIndex:  {status['indexed']} docs  •  "
          f"{status['sections']} sections  •  "
          f"{status['chunks']} chunks  •  "
          f"{status['sentences']} sentences")